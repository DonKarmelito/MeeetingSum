import streamlit as st
import openai
import tempfile
import os
from dotenv import load_dotenv, dotenv_values
from io import BytesIO
from docx import Document
from datetime import datetime
from qdrant_client import QdrantClient
from qdrant_client.models import VectorParams, Distance, PointStruct
import uuid

load_dotenv()
env = dotenv_values(".env")

def get_secret(key):
    """Czyta sekret z .env (lokalnie) lub st.secrets (Streamlit Cloud)."""
    if key in env:
        return env[key]
    if key in st.secrets:
        return st.secrets[key]
    return None

st.set_page_config(
    page_title="MeeetingSum",
    page_icon="uplogo.png",  # plik w folderze projektu
    layout="centered"
)

# ---------- API ----------
if not st.session_state.get("openai_api_key"):
    api_key = get_secret("OPENAI_API_KEY")
    if api_key:
        st.session_state["openai_api_key"] = api_key
    else:
        st.info("Dodaj klucz API OpenAI")
        st.session_state["openai_api_key"] = st.text_input("Klucz API", type="password")
        if st.session_state["openai_api_key"]:
            st.rerun()

if not st.session_state.get("openai_api_key"):
    st.stop()

client = openai.OpenAI(api_key=st.session_state["openai_api_key"])

# ---------- QDRANT ----------
try:
    qdrant = QdrantClient(
        url=get_secret("QDRANT_URL"),
        api_key=get_secret("QDRANT_API_KEY")
    )
    COLLECTION_NAME = "spotkania"
    if not qdrant.collection_exists(COLLECTION_NAME):
        qdrant.create_collection(
            collection_name=COLLECTION_NAME,
            vectors_config=VectorParams(size=1536, distance=Distance.COSINE)
        )
    qdrant_ok = True
except Exception as e:
    st.warning(f"Brak połączenia z Qdrant: {e}")
    qdrant_ok = False

MAX_MB = 25

# ---------- FUNKCJE ----------

def get_embedding(text):
    res = client.embeddings.create(
        model="text-embedding-3-small",
        input=text
    )
    return res.data[0].embedding

def wyodrebnij_audio(plik_video):
    """Wyciąga audio z video i zwraca ścieżkę do pliku mp3."""
    suffix = os.path.splitext(plik_video)[1]
    
    with tempfile.NamedTemporaryFile(suffix=".mp3", delete=False) as tmp_audio:
        tmp_audio_path = tmp_audio.name

    wynik = os.system(
        f'ffmpeg -i "{plik_video}" -vn -acodec libmp3lame -q:a 4 "{tmp_audio_path}" -y -loglevel quiet'
    )
    
    if wynik != 0:
        raise Exception("FFmpeg nie mógł przetworzyć pliku video.")
    
    return tmp_audio_path

def zapisz_do_qdrant(podsumowanie, transkrypcja, nazwa, podsumowanie_en=""):
    embedding = get_embedding(podsumowanie + "\n" + transkrypcja)
    qdrant.upsert(
        collection_name=COLLECTION_NAME,
        points=[
            PointStruct(
                id=str(uuid.uuid4()),
                vector=embedding,
                payload={
                    "podsumowanie": podsumowanie,
                    "podsumowanie_en": podsumowanie_en,
                    "transkrypcja": transkrypcja,
                    "plik": nazwa,
                    "data": datetime.now().isoformat()
                }
            )
        ]
    )

def zapisz_do_pliku(podsumowanie, transkrypcja, nazwa, podsumowanie_en):
    teraz = datetime.now().strftime("%Y-%m-%d %H:%M")

    folder = "notatki"
    os.makedirs(folder, exist_ok=True)

    plik_pl = os.path.join(folder, "notatki_PL.docx")
    plik_en = os.path.join(folder, "notatki_EN.docx")

    if os.path.exists(plik_pl):
        doc_pl = Document(plik_pl)
    else:
        doc_pl = Document()
        doc_pl.add_heading("Notatnik Spotkań", 0)

    doc_pl.add_paragraph("─" * 60)
    doc_pl.add_heading(f"Spotkanie — {teraz}", level=1)
    doc_pl.add_paragraph(f"Plik: {nazwa}")
    doc_pl.add_heading("Podsumowanie", level=2)
    doc_pl.add_paragraph(podsumowanie)
    doc_pl.add_heading("Transkrypcja", level=2)
    doc_pl.add_paragraph(transkrypcja)
    doc_pl.save(plik_pl)

    if os.path.exists(plik_en):
        doc_en = Document(plik_en)
    else:
        doc_en = Document()
        doc_en.add_heading("Meeting Notes", 0)

    doc_en.add_paragraph("─" * 60)
    doc_en.add_heading(f"Meeting — {teraz}", level=1)
    doc_en.add_paragraph(f"File: {nazwa}")
    doc_en.add_heading("Summary", level=2)
    doc_en.add_paragraph(podsumowanie_en)
    doc_en.save(plik_en)

def generuj_docx_pl(podsumowanie, transkrypcja, nazwa):
    doc = Document()
    doc.add_heading("Notatka ze spotkania", 0)
    doc.add_heading("Podsumowanie", level=1)
    doc.add_paragraph(podsumowanie)
    doc.add_heading("Transkrypcja", level=1)
    doc.add_paragraph(transkrypcja)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def generuj_docx_en(podsumowanie_en):
    doc = Document()
    doc.add_heading("Meeting Summary", 0)
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(podsumowanie_en)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def pobierz_wszystkie_jako_docx(jezyk="PL"):
    wyniki = qdrant.scroll(
        collection_name=COLLECTION_NAME,
        limit=100,
        with_payload=True,
        with_vectors=False
    )
    notatki = sorted(
        wyniki[0],
        key=lambda x: x.payload.get("data", ""),
        reverse=True
    )

    doc = Document()

    if jezyk == "PL":
        doc.add_heading("Notatnik Spotkań", 0)
    else:
        doc.add_heading("Meeting Notes", 0)

    for n in notatki:
        p = n.payload
        data = p.get("data", "")[:16]
        plik = p.get("plik", "brak nazwy")

        doc.add_paragraph("─" * 60)

        if jezyk == "PL":
            doc.add_heading(f"Spotkanie — {data}", level=1)
            doc.add_paragraph(f"Plik: {plik}")
            doc.add_heading("Podsumowanie", level=2)
            doc.add_paragraph(p.get("podsumowanie", "brak podsumowania"))
            doc.add_heading("Transkrypcja", level=2)
            doc.add_paragraph(p.get("transkrypcja", "brak transkrypcji"))
        else:
            doc.add_heading(f"Meeting — {data}", level=1)
            doc.add_paragraph(f"File: {plik}")
            doc.add_heading("Summary", level=2)
            # EN summary może nie być w Qdrant - fallback na PL
            doc.add_paragraph(p.get("podsumowanie_en", p.get("podsumowanie", "no summary")))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def reset_sesji():
    st.session_state.transkrypcja = None
    st.session_state.podsumowanie = None
    st.session_state.podsumowanie_en = None
    st.session_state.zapisano_qdrant = False

# ---------- SESSION ----------
if "transkrypcja" not in st.session_state:
    st.session_state.transkrypcja = None
if "podsumowanie" not in st.session_state:
    st.session_state.podsumowanie = None
if "podsumowanie_en" not in st.session_state:
    st.session_state.podsumowanie_en = None
if "nazwa_pliku" not in st.session_state:
    st.session_state.nazwa_pliku = None
if "zapisano_qdrant" not in st.session_state:
    st.session_state.zapisano_qdrant = False
if "poprzedni_plik" not in st.session_state:
    st.session_state.poprzedni_plik = None
if "poprzednie_audio_len" not in st.session_state:
    st.session_state.poprzednie_audio_len = 0


col_tytul, col_logo = st.columns([3, 1])

with col_tytul:
    st.markdown("""
        <style>
            .tytul {
                font-size: 55px;
                font-weight: bold;
                font-family: 'Verdana', sans-serif;
                color: #2C3E50;
                line-height: 1.1;
                margin-bottom: 0px;
            }
            .podtytul {
                font-size: 22px;
                font-weight: normal;
                font-family: 'Verdana', sans-serif;
                color: #5D6D7E;
                margin-top: 0px;
                margin-bottom: 4px;
            }
            .powered {
                font-size: 14px;
                font-family: 'Verdana', sans-serif;
                color: #AAB7B8;
                margin-top: 0px;
            }
        </style>
        <p class="tytul">MeeetingSum</p>
        <p class="podtytul">Podsumowanie spotkań AI</p>
        <p class="powered">Powered by KarmelCodeLab</p>
    """, unsafe_allow_html=True)

with col_logo:
    st.image("logo.png", width=150)



tab1, tab2 = st.tabs(["Audio / Video", "Baza spotkań"])

# ====================================================
# AUDIO
# ====================================================

with tab1:

    upload, record = st.tabs(["Wgraj plik", "Nagraj"])

    # ---------- UPLOAD ----------
    with upload:

        uploaded_file = st.file_uploader(
            "Wgraj plik audio lub video",
            type=["mp3", "wav", "m4a", "mp4", "webm"]
        )

        if uploaded_file:

            if uploaded_file.name != st.session_state.poprzedni_plik:
                reset_sesji()
                st.session_state.poprzedni_plik = uploaded_file.name

            rozmiar_mb = uploaded_file.size / (1024 * 1024)

            if rozmiar_mb > MAX_MB:
                st.error(f"Plik za duży ({rozmiar_mb:.1f}MB). Maksimum to {MAX_MB}MB.")
            else:
                st.caption(f"Rozmiar: {rozmiar_mb:.1f}MB")
                st.audio(uploaded_file)

                if st.button("Transkrybuj plik", key="transkrybuj_plik"):
                    try:
                        jest_video = uploaded_file.name.lower().endswith((".mp4", ".webm", ".mov", ".avi"))

                        with st.spinner("Zapisuję plik..."):
                            with tempfile.NamedTemporaryFile(
                                suffix=os.path.splitext(uploaded_file.name)[1],
                                delete=False
                            ) as tmp:
                                tmp.write(uploaded_file.read())
                                tmp_path = tmp.name

                        # Ekstrakcja audio z video
                        if jest_video:
                            with st.spinner("Wyodrębniam audio z video (może chwilę potrwać)..."):
                                try:
                                    audio_path = wyodrebnij_audio(tmp_path)
                                    os.unlink(tmp_path)  # usuń oryginalny video
                                    tmp_path = audio_path
                                    st.info("✅ Audio wyodrębnione — rozpoczynam transkrypcję...")
                                except Exception as e:
                                    os.unlink(tmp_path)
                                    st.error(f"Błąd ekstrakcji audio: {e}")
                                    st.stop()

                        # Sprawdź rozmiar po ekstrakcji
                        rozmiar_po = os.path.getsize(tmp_path) / (1024 * 1024)
                        if rozmiar_po > MAX_MB:
                            os.unlink(tmp_path)
                            st.error(f"Audio po ekstrakcji nadal za duże ({rozmiar_po:.1f}MB). Maksimum to {MAX_MB}MB.")
                            st.stop()

                        with st.spinner("Transkrybuję..."):
                            with open(tmp_path, "rb") as f:
                                transcript = client.audio.transcriptions.create(
                                    model="whisper-1",
                                    file=f
                                )
                            os.unlink(tmp_path)

                        st.session_state.transkrypcja = transcript.text
                        st.session_state.nazwa_pliku = uploaded_file.name

                    except Exception as e:
                        st.error(f"Błąd transkrypcji: {e}")

    # ---------- RECORD ----------
    with record:
        st.write("Nagrywanie w przeglądarce")

        # natywny rejestrator Streamlit
        audio_file = st.audio_input("Nagraj notatkę")

        if audio_file is None:
            st.info("Brak nagrania – kliknij mikrofon, powiedz kilka sekund i zatrzymaj.")
            st.stop()

        # podgląd
        st.audio(audio_file)

        # wyślij do Whisper
        if st.button("Transkrybuj nagranie", key="transkrybuj_nagranie_webrtc"):
            try:
                with st.spinner("Transkrybuję..."):
                    audio_file.seek(0)
                    # Streamlit zwykle daje WAV; nazwijmy plik .wav
                    transcript = client.audio.transcriptions.create(
                        model="whisper-1",
                        file=("nagranie.wav", audio_file, "audio/wav")
                    )
                st.session_state.transkrypcja = transcript.text
                st.session_state.nazwa_pliku = "nagranie"
                st.success("Transkrypcja gotowa.")
            except Exception as e:
                st.error(f"Błąd transkrypcji: {e}")

    # ====================================================
    # TRANSKRYPCJA I PODSUMOWANIE
    # ====================================================

    if st.session_state.transkrypcja:

        with st.expander("Pokaż transkrypcję"):
            st.text_area("Transkrypcja", st.session_state.transkrypcja, height=300)

        st.caption(f"Długość: {len(st.session_state.transkrypcja)} znaków")

        if st.button("Generuj podsumowanie", key="generuj_podsumowanie"):
            try:
                with st.spinner("Generuję podsumowanie..."):
                    response = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[
                            {
                                "role": "system",
                                "content": "Stwórz krótkie podsumowanie spotkania. Wypisz tematy, decyzje i zadania."
                            },
                            {
                                "role": "user",
                                "content": st.session_state.transkrypcja
                            }
                        ]
                    )
                st.session_state.podsumowanie = response.choices[0].message.content

                with st.spinner("Tłumaczę na angielski..."):
                    response_en = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[
                            {
                                "role": "system",
                                "content": "Translate this meeting summary to professional English."
                            },
                            {
                                "role": "user",
                                "content": st.session_state.podsumowanie
                            }
                        ]
                    )
                st.session_state.podsumowanie_en = response_en.choices[0].message.content
                st.session_state.zapisano_qdrant = False

            except Exception as e:
                st.error(f"Błąd generowania podsumowania: {e}")

    if st.session_state.podsumowanie:

        st.subheader("Podsumowanie (PL)")
        st.write(st.session_state.podsumowanie)

        st.subheader("Summary (EN)")
        st.write(st.session_state.podsumowanie_en)

        teraz = datetime.now().strftime("%Y-%m-%d_%H-%M")

        col1, col2, col3, col4 = st.columns(4)

        with col1:
            if not qdrant_ok:
                st.warning("Brak połączenia z bazą")
            elif st.session_state.zapisano_qdrant:
                st.success("Zapisano w bazie!")
            else:
                if st.button("Zapisz w bazie", key="zapisz_baza_main"):
                    try:
                        with st.spinner("Zapisuję..."):
                            zapisz_do_qdrant(
                                st.session_state.podsumowanie,
                                st.session_state.transkrypcja,
                                st.session_state.nazwa_pliku,
                                st.session_state.podsumowanie_en
                            )
                        st.session_state.zapisano_qdrant = True
                        st.rerun()
                    except Exception as e:
                        st.error(f"Błąd zapisu do bazy: {e}")

        with col2:
            buf_pl = generuj_docx_pl(
                st.session_state.podsumowanie,
                st.session_state.transkrypcja,
                st.session_state.nazwa_pliku
            )
            st.download_button(
                "Pobierz PL",
                buf_pl,
                file_name=f"spotkanie_{teraz}_PL.docx",
                key="pobierz_pl"
            )

        with col3:
            buf_en = generuj_docx_en(st.session_state.podsumowanie_en)
            st.download_button(
                "Pobierz EN",
                buf_en,
                file_name=f"meeting_{teraz}_EN.docx",
                key="pobierz_en"
            )

        with col4:
            if not qdrant_ok:
                st.warning("Brak połączenia z bazą")
            elif st.session_state.zapisano_qdrant:
                st.success("Zapisano!")
            else:
                if st.button("Zapisz w bazie", key="zapisz_baza"):
                    try:
                        with st.spinner("Zapisuję..."):
                            zapisz_do_qdrant(
                                st.session_state.podsumowanie,
                                st.session_state.transkrypcja,
                                st.session_state.nazwa_pliku,
                                st.session_state.podsumowanie_en
                            )
                        st.session_state.zapisano_qdrant = True
                        st.rerun()
                    except Exception as e:
                        st.error(f"Błąd zapisu do bazy: {e}")

# ====================================================
# BAZA
# ====================================================

with tab2:

    st.subheader("Wyszukaj spotkania")

    if qdrant_ok:
        teraz = datetime.now().strftime("%Y-%m-%d_%H-%M")
        col_dl1, col_dl2 = st.columns(2)

        with col_dl1:
            try:
                buf_all_pl = pobierz_wszystkie_jako_docx(jezyk="PL")
                st.download_button(
                    "📥 Pobierz wszystkie (PL)",
                    buf_all_pl,
                    file_name=f"wszystkie_spotkania_PL_{teraz}.docx",
                    key="pobierz_wszystkie_pl"
                )
            except Exception as e:
                st.error(f"Błąd generowania pliku: {e}")

        with col_dl2:
            try:
                buf_all_en = pobierz_wszystkie_jako_docx(jezyk="EN")
                st.download_button(
                    "📥 Pobierz wszystkie (EN)",
                    buf_all_en,
                    file_name=f"all_meetings_EN_{teraz}.docx",
                    key="pobierz_wszystkie_en"
                )
            except Exception as e:
                st.error(f"Błąd generowania pliku: {e}")

    if not qdrant_ok:
        st.error("Brak połączenia z Qdrant.")
    else:
        query = st.text_input("Wpisz temat")

        if query:
            try:
                with st.spinner("Szukam..."):
                    emb = get_embedding(query)
                    try:
                        # nowe API qdrant-client (>=1.7)
                        hits = qdrant.query_points(
                            collection_name=COLLECTION_NAME,
                            query=emb,
                            limit=5,
                            with_payload=True
                        ).points
                    except AttributeError:
                        # fallback na starsze wersje
                        hits = qdrant.search(
                            collection_name=COLLECTION_NAME,
                            query_vector=emb,
                            limit=5,
                            with_payload=True
                        )

                if hits:
                    for h in hits:
                        p = h.payload
                        trafnosc = round(h.score * 100, 1)
                        data = p.get("data", "brak daty")[:16]
                        plik = p.get("plik", "brak nazwy")
                        with st.expander(f"{data} — {plik} (trafność: {trafnosc}%)"):
                            st.write(p.get("podsumowanie", "brak podsumowania"))
                else:
                    st.info("Brak wyników.")
            except Exception as e:
                st.error(f"Błąd wyszukiwania: {e}")


        st.divider()
        st.subheader("Wszystkie spotkania")

        if st.button("Załaduj wszystkie", key="zaladuj_wszystkie"):
            try:
                wyniki = qdrant.scroll(
                    collection_name=COLLECTION_NAME,
                    limit=100,
                    with_payload=True,
                    with_vectors=False
                )
                notatki = sorted(
                    wyniki[0],
                    key=lambda x: x.payload.get("data", ""),
                    reverse=True
                )

                if notatki:
                    for n in notatki:
                        p = n.payload
                        data = p.get("data", "brak daty")[:16]
                        plik = p.get("plik", "brak nazwy")
                        with st.expander(f"{data} — {plik}"):
                            st.write(p.get("podsumowanie", "brak podsumowania"))
                else:
                    st.info("Brak zapisanych spotkań.")

            except Exception as e:
                st.error(f"Błąd ładowania: {e}")