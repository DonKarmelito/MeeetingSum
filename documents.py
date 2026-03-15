from io import BytesIO
from docx import Document


def generuj_docx_pl(podsumowanie, transkrypcja, nazwa):
    """Generuje plik Word z podsumowaniem i transkrypcją po polsku."""
    doc = Document()
    doc.add_heading("Notatka ze spotkania", 0)
    doc.add_heading("Podsumowanie", level=1)
    doc.add_paragraph(podsumowanie)
    doc.add_heading("Transkrypcja", level=1)
    doc.add_paragraph(transkrypcja)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generuj_docx_en(podsumowanie_en):
    """Generuje plik Word z podsumowaniem po angielsku."""
    doc = Document()
    doc.add_heading("Meeting Summary", 0)
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(podsumowanie_en)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def pobierz_wszystkie_jako_docx(qdrant, collection_name, jezyk="PL"):
    """
    Generuje plik Word ze wszystkimi spotkaniami z bazy Qdrant.
    jezyk: "PL" lub "EN"
    """
    wyniki = qdrant.scroll(
        collection_name=collection_name,
        limit=100,
        with_payload=True,
        with_vectors=False
    )
    notatki = sorted(
        wyniki[0],
        key=lambda x: x.payload.get("data", ""),
        reverse=True
    )

    doc = Document()

    if jezyk == "PL":
        doc.add_heading("Notatnik Spotkań", 0)
    else:
        doc.add_heading("Meeting Notes", 0)

    for n in notatki:
        p = n.payload
        data = p.get("data", "")[:16]
        plik = p.get("plik", "brak nazwy")

        doc.add_paragraph("─" * 60)

        if jezyk == "PL":
            doc.add_heading(f"Spotkanie — {data}", level=1)
            doc.add_paragraph(f"Plik: {plik}")
            doc.add_heading("Podsumowanie", level=2)
            doc.add_paragraph(p.get("podsumowanie", "brak podsumowania"))
            doc.add_heading("Transkrypcja", level=2)
            doc.add_paragraph(p.get("transkrypcja", "brak transkrypcji"))
        else:
            doc.add_heading(f"Meeting — {data}", level=1)
            doc.add_paragraph(f"File: {plik}")
            doc.add_heading("Summary", level=2)
            doc.add_paragraph(p.get("podsumowanie_en", p.get("podsumowanie", "no summary")))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
